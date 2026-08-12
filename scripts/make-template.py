#!/usr/bin/env python3
"""
Generates the Word content template Drishti fills in and hands back.

Design intent: the five case studies are already written and live. This form
does NOT ask her to rewrite them. It collects exactly three things:
  1. Facts only she has (the 23 visible placeholders on the site).
  2. Corrections to the drafts I wrote from her resume.
  3. The image files, with specs and a clearance check per shot.

Run:  python3 scripts/make-template.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT

GREY = RGBColor(0x66, 0x66, 0x66)
ACCENT = RGBColor(0xA8, 0x3D, 0x24)

# ---------------------------------------------------------------- helpers


def h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def para(doc, text="", *, italic=False, grey=False, size=10.5, space_after=6, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    r.font.size = Pt(size)
    if grey:
        r.font.color.rgb = GREY
    return p


def label(doc, text):
    """Small accent eyebrow."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = ACCENT
    return p


def draft(doc, text):
    """My draft, clearly marked as editable."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("DRAFT  ")
    r.bold = True
    r.font.size = Pt(7.5)
    r.font.color.rgb = GREY
    r2 = p.add_run(text)
    r2.italic = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = GREY
    return p


def answer(doc, hint="", lines=1):
    """A highlighted blank for her to type into."""
    for _ in range(lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(hint if hint else " " * 90)
        r.font.size = Pt(10.5)
        r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return p


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0].cells
    for i, name in enumerate(headers):
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(name)
        r.bold = True
        r.font.size = Pt(8.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            if val == "__FILL__":
                r = p.add_run("            ")
                r.font.highlight_color = WD_COLOR_INDEX.YELLOW
                r.font.size = Pt(10)
            else:
                r = p.add_run(str(val))
                r.font.size = Pt(9.5)
                if str(val).startswith("["):
                    r.italic = True
                    r.font.color.rgb = GREY
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("─" * 62)
    r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    r.font.size = Pt(8)


# ---------------------------------------------------------------- content

STUDIES = [
    {
        "n": "01",
        "slug": "agentic-support",
        "title": "Deciding what the agent does alone",
        "org": "Autodesk",
        "role": "Lead product designer — in-product agentic assistance",
        "timeframe": "2024 – 2026",
        "kind": "Deep case study (hero)",
        "brief": {
            "situation": "Autodesk wanted in-product AI assistance across three technical products. Engineering wanted a model-confidence threshold to decide when the agent could act alone; support operations wanted approval on everything.",
            "decision": "I refused the binary and owned the autonomy policy instead: a four-rung ladder (answer / suggest / stage / act), with the rung gated on consequence-of-being-wrong rather than model confidence alone.",
            "defend": "The escalation redesign, which we isolated against a holdback. Consequence is classified once by humans in the action registry, never inferred by the model at runtime.",
            "wrong": "I owned rung assignment centrally and became the bottleneck across three product teams. And I raised instrumentation too late to cleanly attribute the headline 91%.",
        },
        "metrics": [
            ("91%", "Instant digital resolution"),
            ("60%", "Improvement in support effectiveness"),
            ("3", "Products, one autonomy policy"),
        ],
        "blanks": [
            "What confidence threshold did engineering actually propose? (e.g. “85%”, or “they never named a number”)",
            "How many support transcripts did you analyse in discovery?",
            "How many sessions did you run with support agents and customers?",
            "How long before central rung-assignment made you the bottleneck? (e.g. “about two months”)",
            "Which product did the escalation holdback run on? (AutoCAD / Revit / Fusion)",
            "How much did human-agent handling time move on escalated conversations? (the number you CAN defend)",
            "When did you raise instrumentation, and what stage was the rollout at by then?",
            "Is Autodesk Assistant publicly reachable? Paste the URL, or write “internal only”.",
        ],
        "shots": [
            ("HIGHEST VALUE — the staged-action review UI in Revit: the agent's proposed change as a reviewable diff, affected model elements highlighted, single-step reject visible.", "1600×1000 or wider", "Redact model/customer names"),
            ("The escalation handoff as the human agent sees it — the summary, what was ruled out, product state.", "1600×1000", "Redact customer data"),
            ("The autonomy ladder as it exists in your own docs, if you have a version.", "any", "Check internal-doc policy"),
        ],
    },
    {
        "n": "02",
        "slug": "customer-success-planning",
        "title": "A tool for a job nobody had written down",
        "org": "Autodesk",
        "role": "Lead product designer — end to end",
        "timeframe": "2024 – 2025 · launched July 2025",
        "kind": "Deep case study",
        "brief": {
            "situation": "“Build a customer success planning tool” arrived with executive backing and no definition of a success plan. Six stakeholders gave six different definitions.",
            "decision": "I changed the brief. Six weeks of JTBD discovery showed the plan was never the point — the alignment conversation was. I brought that back as a fork with costs attached, not a verdict, and the room picked it.",
            "defend": "The revisit rate — plans updated rather than written once. It is the behaviour the reframe predicted and it needs no control group.",
            "wrong": "I cut customer-side notifications to protect scope, which removed the only support for the weakest link in the loop. And my storyboards persuaded leadership while under-specifying the build.",
        },
        "metrics": [
            ("71.8%", "Retention, first year post-launch"),
            ("19.2%", "Renewal rate, same cohort"),
            ("Jul 2025", "Shipped"),
        ],
        "blanks": [
            "How many interviews with CSMs and account teams?",
            "How many mid-market CUSTOMERS did you interview? (this is the number that makes the reframe credible)",
            "When was engineering blocked on decisions you hadn't made yet? (roughly which month/phase)",
            "What was the plan revisit rate? (the metric you said you'd keep over retention)",
            "Can any of this be shown or linked — a redacted walkthrough, a demo? Or “internal only”.",
        ],
        "shots": [
            ("The workshop wall — the “what the customer would say in twelve months” sentences, clustered. A phone photo of sticky notes is genuinely fine here.", "1600×900", "Redact company + customer names"),
            ("Two or three storyboard frames from the twelve-month arc — ideally the joint-definition moment and the quarterly revisit.", "1600×900", "Low-fidelity is better here"),
            ("The shipped joint-definition screen.", "1600×1000", "Redact all customer data"),
        ],
    },
    {
        "n": "03",
        "slug": "account-experience",
        "title": "Six weeks to a shared picture",
        "org": "Autodesk",
        "role": "Product designer — vision, then execution",
        "timeframe": "2023 – present",
        "kind": "Deep case study",
        "brief": {
            "situation": "Everyone agreed the signed-in account experience was broken; nobody agreed what fixed looked like. Billing, SMB, web, and support each had a different, defensible roadmap.",
            "decision": "I built the artifact that resolved the disagreement rather than a comp set: three horizons, each with a stated bet and a stated falsifier — which turned a vision presentation into a funding conversation.",
            "defend": "Senior stakeholders could restate the direction accurately in meetings I was not in. That was the goal, and it held.",
            "wrong": "I treated the vision as a deliverable with an end date. It decayed as people changed roles, and the falsifiers were written but not all instrumented.",
        },
        "metrics": [
            ("6 weeks", "Ambiguity to executive-ready vision"),
            ("Roadmap", "Vision shaped the long-term account plan"),
            ("Ongoing", "Experiment sequence in the SMB portal"),
        ],
        "blanks": [
            "Subscription-state legibility experiment — what was the measured result?",
            "Renewal price-change comprehension — result, or write “inconclusive” (that is a fine answer and I will print it)",
            "Seat assignment for very small teams — what date/quarter is it in flight as of?",
            "How long after the vision landed did it start decaying? (e.g. “about nine months”)",
            "Is the account portal publicly reachable? Paste the URL, or “sign-in required”.",
        ],
        "shots": [
            ("Two or three frames from the vision narrative — grey-box concepts with real customer language. The “what do I own” frame is the strongest single image.", "1600×900", "Redact account data"),
            ("The horizon/falsifier artifact as it actually existed.", "1600×900", "Check if pre-decisional"),
            ("A shipped SMB portal screen — ideally subscription state.", "1600×1000", "Redact all account data"),
        ],
    },
    {
        "n": "04",
        "slug": "commercial-banking-portal",
        "title": "Trust in a financial workflow",
        "org": "American Express (via IntraEdge)",
        "role": "UX designer — Global Commercial Services",
        "timeframe": "2021 – 2022",
        "kind": "Short piece",
        "brief": {
            "situation": "Corporate clients reconciling statements and pulling reports other people make decisions on. Arriving with a specific question, under time pressure, accountable for the number.",
            "decision": "I dropped the consumer instinct to simplify and consistently surfaced provenance instead — what a figure covers, when it was calculated, what it excludes.",
            "defend": "That principle. It transferred directly to designing AI experiences: an agent's answer and a statement figure have the same requirement.",
            "wrong": "I optimised screens well and never mapped the reconciliation job end to end. I also accepted the metrics I was handed without asking how they were constructed.",
        },
        "metrics": [
            ("+30%", "Increase in user engagement"),
            ("−40%", "Reduction in identified pain points"),
        ],
        "blanks": [
            "Anything you can show at all from this work? Amex client data is the tightest constraint on the site — “nothing publishable” is a completely acceptable answer.",
        ],
        "shots": [
            ("The client-access dashboard, ideally a report state showing provenance (date range, last updated, exclusions). ONLY if clearance allows.", "1600×1000", "CLEAR WITH AMEX / INTRAEDGE FIRST"),
        ],
    },
    {
        "n": "05",
        "slug": "working-with-ai",
        "title": "How I actually work with AI",
        "org": "Practice note",
        "role": "—",
        "timeframe": "Current",
        "kind": "Short piece",
        "brief": {
            "situation": "Every team asks how a designer uses AI, and most answers are either defensive or credulous.",
            "decision": "I keep a working ledger: I hand over work where I can evaluate output faster than I could produce it, and keep work where evaluating the output is the skill.",
            "defend": "Reversibility is the price of autonomy. Confidence is a property of the model; consequence is a property of the world.",
            "wrong": "I have twice gotten attached to a prototype because I built it. Sunk cost applies to your own code with surprising force.",
        },
        "metrics": [
            ("HTML/CSS/JS", "Prototypes I build myself, AI-assisted"),
            ("200+", "Attendees, Autodesk TechX 2026 talk"),
        ],
        "blanks": [
            "Which project did you over-defend a prototype on, and for roughly how long?",
            "Can the TechX talk recording or deck be shared publicly? Paste the link — this is the single highest-leverage asset you're not using.",
            "Is the AI do/don't ledger accurate? Anything to add or remove? (see the live page)",
        ],
        "shots": [
            ("A photo or still of you presenting at TechX — even a phone shot from the audience.", "1600×1000", "Check event photo policy"),
            ("A screenshot of one of your own code prototypes running.", "1600×1000", "Redact any real data"),
        ],
    },
]


def build():
    doc = Document()

    # Base style
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)

    for s in doc.sections:
        s.left_margin = Inches(0.9)
        s.right_margin = Inches(0.9)
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)

    # ------------------------------------------------ cover
    t = doc.add_heading("Portfolio content — what I need from you", level=0)
    para(
        doc,
        "Five case studies are already written and live on the site. This form does not ask you "
        "to rewrite them. It collects the three things I cannot get without you.",
        size=11,
    )

    label(doc, "How to use this")
    for i, line in enumerate(
        [
            "Anything in YELLOW is a blank for you to type into.",
            "Anything in GREY ITALIC marked DRAFT is what I wrote from your résumé. Read it, and correct it. "
            "If a draft is right, write “ok” next to it — that is a valid answer and saves you time.",
            "Skip anything you genuinely don't have. A missing number costs you far less than a number you "
            "can't source in an interview. Where you skip, I cut the sentence rather than soften it.",
            "Do NOT round numbers up to sound better. Every metric on this site is published with its caveat, "
            "which is what makes the rest of them believable.",
            "Images: don't paste them into this document. Save them in a folder, name them as shown in each "
            "shot list, and send the folder.",
        ]
    ):
        p = doc.add_paragraph(line, style="List Number")
        p.paragraph_format.space_after = Pt(5)
        for r in p.runs:
            r.font.size = Pt(10)

    label(doc, "The one thing that matters most")
    para(
        doc,
        "Hiring managers screen portfolios by asking “is it live and traceable?” before anything else. "
        "Every case study below has a “can this be linked or shown” question. Those five answers are worth "
        "more than all the prose corrections combined.",
        size=10.5,
    )

    doc.add_page_break()

    # ------------------------------------------------ site-wide
    h(doc, "Part 1 — Site-wide", level=1)

    label(doc, "1.1  Two facts to settle")
    table(
        doc,
        ["Question", "Why it matters", "Your answer"],
        [
            ["Five years or six?", "Your résumé says “5+”, the site says six. They must match.", "__FILL__"],
            ["Domain you'll publish on", "Sets canonical URLs, sitemap, and structured data.", "__FILL__"],
        ],
        widths=[1.9, 2.7, 1.9],
    )

    label(doc, "1.2  Mentoring — the weakest evidence on the site")
    para(
        doc,
        "This row is currently a visible placeholder on the Approach page, and a Staff interviewer will "
        "probe it. Write real specifics or tell me to delete the row — thin evidence is worse than none.",
        size=10,
        grey=True,
        italic=True,
    )
    table(
        doc,
        ["How many designers", "Over what period", "Formal or informal", "One concrete thing that changed for them"],
        [["__FILL__", "__FILL__", "__FILL__", "__FILL__"]],
        widths=[1.2, 1.2, 1.2, 2.9],
    )

    label(doc, "1.3  Have you shared your prototyping workflow with other designers?")
    para(
        doc,
        "A lunch-and-learn, a template, a written guide, onboarding a teammate onto Claude Code / Cursor. "
        "Practice improvement is a Staff-level signal and this is your strongest available example.",
        size=10,
        grey=True,
        italic=True,
    )
    answer(doc, lines=2)

    label(doc, "1.4  Links I should add")
    table(
        doc,
        ["Asset", "Link or “no”"],
        [
            ["TechX 2026 talk recording or deck", "__FILL__"],
            ["Autodesk University workshop material", "__FILL__"],
            ["Anything else public with your name on it", "__FILL__"],
        ],
        widths=[3.6, 2.9],
    )

    doc.add_page_break()

    # ------------------------------------------------ per study
    h(doc, "Part 2 — The five case studies", level=1)
    para(
        doc,
        "One form per study. Read the live page first (localhost:3000/work/<slug>) — it is faster to "
        "correct than to describe from memory.",
        grey=True,
        italic=True,
        size=10,
    )

    for s in STUDIES:
        doc.add_page_break()
        h(doc, f"{s['n']}  {s['title']}", level=1)
        para(doc, f"{s['kind']}  ·  /work/{s['slug']}", grey=True, italic=True, size=9.5)

        # A. basics
        label(doc, "A.  Basics — correct anything wrong")
        table(
            doc,
            ["Field", "My draft", "Correction (leave blank if right)"],
            [
                ["Title", s["title"], "__FILL__"],
                ["Organisation", s["org"], "__FILL__"],
                ["Your role", s["role"], "__FILL__"],
                ["Timeframe", s["timeframe"], "__FILL__"],
            ],
            widths=[1.0, 3.0, 2.5],
        )

        label(doc, "A2.  Who else was on it, and what was exclusively yours")
        para(
            doc,
            "Interviewers ask this in the first two minutes. Being specific about what was NOT yours makes "
            "the rest of the claim credible.",
            grey=True,
            italic=True,
            size=9.5,
        )
        table(
            doc,
            ["The team / functions involved", "What was exclusively yours", "What was NOT yours"],
            [["__FILL__", "__FILL__", "__FILL__"]],
            widths=[2.2, 2.2, 2.1],
        )

        # B. brief
        label(doc, "B.  The 60-second brief — this is the block at the top of the page")
        for key, name in [
            ("situation", "The situation"),
            ("decision", "The decision you owned"),
            ("defend", "What you'd defend"),
            ("wrong", "What you got wrong"),
        ]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(8)
            r = p.add_run(name)
            r.bold = True
            r.font.size = Pt(10)
            draft(doc, s["brief"][key])
            answer(doc, "→ ", lines=1)

        # C. metrics
        doc.add_page_break()
        label(doc, "C.  Metrics — the part interviewers attack")
        para(
            doc,
            "For each number: what was it before, how was it measured, and what ELSE changed in the same "
            "window. If you can't answer the middle column, say so and I'll present it as unverifiable.",
            grey=True,
            italic=True,
            size=9.5,
        )
        rows = []
        for val, lab in s["metrics"]:
            rows.append([f"{val}\n{lab}", "__FILL__", "__FILL__", "__FILL__"])
        table(
            doc,
            ["The number", "Baseline before it", "How it was measured / who owns it", "What else changed in the same window"],
            rows,
            widths=[1.5, 1.5, 1.9, 1.7],
        )
        label(doc, "C2.  Which single number here would you defend hardest under scrutiny? And which least?")
        answer(doc, "→ ", lines=2)

        # D. live
        label(doc, "D.  Live and traceable — highest-value question on this form")
        table(
            doc,
            ["Question", "Your answer"],
            [
                ["Public URL, if any", "__FILL__"],
                ["If not public — can anything be shown? (redacted screens, a demo video, nothing)", "__FILL__"],
                ["Is it still in production / scaling?", "__FILL__"],
            ],
            widths=[3.9, 2.6],
        )

        # E. corrections
        label(doc, "E.  Read the live page — what's factually wrong or missing?")
        para(
            doc,
            "Don't rewrite it. Just list what's wrong, and anything important I invented or left out.",
            grey=True,
            italic=True,
            size=9.5,
        )
        table(
            doc,
            ["Section / sentence", "What's wrong, or what's missing"],
            [["__FILL__", "__FILL__"], ["__FILL__", "__FILL__"], ["__FILL__", "__FILL__"]],
            widths=[2.4, 4.1],
        )

        # F. blanks
        label(doc, "F.  Specific blanks currently visible on this page")
        para(
            doc,
            "These render as amber highlights on the live site until you fill them.",
            grey=True,
            italic=True,
            size=9.5,
        )
        for q in s["blanks"]:
            p = doc.add_paragraph(q, style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            for r in p.runs:
                r.font.size = Pt(10)
            answer(doc, "→ ", lines=1)

        # G. imagery
        label(doc, "G.  Imagery for this study")
        rows = []
        for i, (desc, spec, clearance) in enumerate(s["shots"], start=1):
            rows.append([f"{s['slug']}-{i}", desc, spec, clearance, "__FILL__"])
        table(
            doc,
            ["File name to use", "What to capture", "Min size", "Clearance check", "Have it? Y/N"],
            rows,
            widths=[1.25, 2.5, 0.9, 1.15, 0.7],
        )

    # ------------------------------------------------ image specs
    doc.add_page_break()
    h(doc, "Part 3 — Image specs", level=1)

    label(doc, "3.1  Technical")
    table(
        doc,
        ["Item", "Spec"],
        [
            ["Format", "PNG for UI screens, JPG for photos. I'll convert and generate WebP."],
            ["Minimum width", "1600px for landscape screens. 1000px for portraits. Bigger is fine — I downscale."],
            ["Aspect", "Don't crop to fit. Send the full frame; I'll crop per slot."],
            ["Retina", "If you can export @2x, do — it costs nothing and I'll use it."],
            ["Naming", "Use the file names in each shot list above (e.g. agentic-support-1.png)."],
            ["Delivery", "One folder, not pasted into this document. Word recompresses images badly."],
        ],
        widths=[1.6, 4.9],
    )

    label(doc, "3.2  Redaction checklist — run this on every screen before sending")
    for line in [
        "Customer and company names removed or replaced with plausible fakes.",
        "Real email addresses, account numbers, and invoice/order IDs removed.",
        "Currency amounts — check whether these are real customer figures.",
        "Internal-only URLs, admin tooling, and ticket IDs cropped out.",
        "Colleagues' names and avatars in comment threads or assignment fields.",
        "For Amex: assume everything is client data until someone tells you otherwise.",
    ]:
        p = doc.add_paragraph(line, style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        for r in p.runs:
            r.font.size = Pt(10)

    label(doc, "3.3  What I do NOT need")
    para(
        doc,
        "No persona sheets, no empathy maps, no Double Diamond diagrams, no redline or spacing specs, "
        "no mood boards. Design hiring managers now read process-framework diagrams as a red flag — a "
        "signal that a methodology is standing in for evidence of judgement. The site deliberately has none, "
        "and every diagram on it is built from your actual decisions.",
        size=10,
    )

    label(doc, "3.4  If you can only produce ONE image")
    para(
        doc,
        "Make it the staged-action review UI in Revit (agentic-support-1). It makes “the agent does the work "
        "but a person commits it” concrete, and that is the single most differentiated idea in your portfolio.",
        size=10,
    )

    out = "Portfolio-Content-Template.docx"
    doc.save(out)
    print(f"saved {out}")


if __name__ == "__main__":
    build()
